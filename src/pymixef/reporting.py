"""Manifest-driven human-readable reports."""

from __future__ import annotations

import html
from pathlib import Path
from typing import Any

from .results import FitResult


def _markdown(result: FitResult) -> str:
    parameter_rows = "\n".join(
        f"| `{name}` | {value:.10g} |" for name, value in result.parameters.items()
    )
    warning_rows = "\n".join(
        f"- `{warning.code}` ({warning.severity}): {warning.message}"
        for warning in (*result.warnings, *result.convergence.warnings)
    )
    if not warning_rows:
        warning_rows = "No structured warnings were emitted."
    covariance_method = result.extra.get("covariance_method", "backend default")
    approximation = result.extra.get("approximation", result.engine)
    return f"""# PyMixEF model result

## 1. Model purpose and analysis population

This report was generated from the immutable result and run manifest. Scientific
purpose and estimand context must be supplied by the analysis owner.

## 2. Data audit

- Analysis observations: {result.n_observations}
- Input data hash: `{result.manifest.data_hash}`
- Model IR hash: `{result.manifest.model_ir_hash}`

## 3. Mathematical model and transformations

The archived model IR is included in the machine-readable result. Parameter values
below are reported on their natural scale unless the model IR states otherwise.

## 4. Estimator and approximation

- Engine: `{result.engine}`
- Method: `{result.method}`
- Approximation: `{approximation}`
- Objective: `{result.objective:.12g}`
- Log likelihood: `{result.log_likelihood}`

## 5. Convergence and numerical diagnostics

- Status: `{result.convergence.status}`
- Optimizer terminated: `{result.convergence.optimizer_terminated}`
- Scaled gradient infinity norm: `{result.convergence.scaled_gradient_inf_norm}`
- Hessian positive definite: `{result.convergence.hessian.positive_definite}`
- Conditional-mode failures: `{result.convergence.conditional_mode_failures}`
- ODE failures: `{result.convergence.ode_failures}`

## 6. Parameter estimates and uncertainty

| Parameter | Estimate |
|---|---:|
{parameter_rows}

Uncertainty method: `{covariance_method}`. See the machine-readable result for the
full covariance matrix and approximation metadata.

## 7. Random effects and predictions

Random-effect estimates and explicit conditional/population prediction arrays are
stored in the result object when supported by the selected backend.

## 8. Diagnostic data

Available serialized tables: {", ".join(sorted(result.diagnostic_data)) or "none"}.
Plots are intentionally a separate presentation layer.

## 9. Sensitivity analyses

No sensitivity analysis is implied by this primary fit. Use approximation,
quadrature, tolerance, bootstrap, or robust-comparison workflows explicitly.

## 10. Simulation

The result declares reproducibility class
`{result.manifest.reproducibility_class}` and records all supplied seeds.

## 11. Known limitations and warnings

{warning_rows}

## 12. Run manifest

- PyMixEF version: `{result.manifest.package_version}`
- Manifest schema: `{result.manifest.manifest_schema_version}`
- Created UTC: `{result.manifest.created_at_utc}`
- Environment and thread settings are present in the archived manifest.

> PyMixEF supplies evidence and reproducibility tooling; this report is not a universal regulatory validation certificate.
"""


def render_report(result: FitResult, path: str | Path) -> Path:
    """Render a result as Markdown, HTML, PDF, or Word.

    PDF and Word require the optional ``report`` dependencies.  All formats are
    generated from the same immutable Markdown content.
    """

    destination = Path(path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    markdown = _markdown(result)
    suffix = destination.suffix.lower()
    if suffix in {".md", ".markdown"}:
        destination.write_text(markdown, encoding="utf-8")
        return destination
    if suffix == ".html":
        try:
            import markdown as markdown_package

            body = markdown_package.markdown(markdown, extensions=["tables"])
        except ImportError:
            body = "<pre>" + html.escape(markdown) + "</pre>"
        destination.write_text(
            "<!doctype html><html lang='en'><meta charset='utf-8'>"
            "<title>PyMixEF report</title><style>"
            "body{font:16px system-ui;max-width:960px;margin:2rem auto;line-height:1.5}"
            "table{border-collapse:collapse}th,td{border:1px solid #bbb;padding:.4rem}"
            "code{background:#f5f5f5;padding:.1rem .25rem}</style>"
            f"<body>{body}</body></html>",
            encoding="utf-8",
        )
        return destination
    if suffix == ".pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as error:
            raise ImportError("PDF reports require `pip install pymixef[report]`.") from error
        styles = getSampleStyleSheet()
        story: list[Any] = []
        for block in markdown.split("\n\n"):
            cleaned = html.escape(block).replace("\n", "<br/>")
            style = styles["Heading1"] if block.startswith("# ") else styles["BodyText"]
            story.extend((Paragraph(cleaned.lstrip("# "), style), Spacer(1, 0.12 * inch)))
        document = SimpleDocTemplate(
            str(destination),
            pagesize=letter,
            rightMargin=0.65 * inch,
            leftMargin=0.65 * inch,
            topMargin=0.65 * inch,
            bottomMargin=0.65 * inch,
        )
        document.build(story)
        return destination
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as error:
            raise ImportError("Word reports require `pip install pymixef[report]`.") from error
        document = Document()
        for line in markdown.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:], level=0)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=1)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            elif line and not line.startswith("|"):
                document.add_paragraph(line)
        document.save(destination)
        return destination
    raise ValueError("Report path must end in .md, .html, .pdf, or .docx.")
